from underthesea import word_tokenize
from nltk.tokenize import sent_tokenize
from gensim.test.utils import datapath
import gensim
import pandas as pd
import random
import string
import pickle
import os
import codecs
import smart_open
import PyPDF2
from docx import Document

# lấy dữ liệu text
def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        text = "\n".join(sent_tokenize(para.text))
        fullText.append(text)
    return '\n'.join(fullText)

# chuyển sang file txt
def change_file_extension (path):
    split_tup = os.path.splitext(path)
    file_extension = split_tup[1]
    file_name = split_tup[0] + ".txt"
    if file_extension == ".pdf":
        pdf = open(path,"rb")
        pdfReader = PyPDF2.PdfFileReader(pdf)
        numOfPages=pdfReader.numPages
        for i in range(numOfPages):
            page = pdfReader.getPage(i)
            text = page.extractText()
            with codecs.open(file_name, 'w', "utf-8-sig") as o_file:
                o_file.write(text)
        pdf.close()
        o_file.close()
    if file_extension == ".docx":
        document = getText(path)
        with codecs.open(file_name, 'w', "utf-8-sig") as o_file:
            o_file.write(document)
        o_file.close()
    return file_name

# tạo corpus và df cho file test
def read_corpus(my_dict, test_file):
    i = 0
    with smart_open.open(test_file, encoding='utf-8-sig') as f:
        for line in (f):
            new_string = line.translate(str.maketrans('', '', string.punctuation))
            tokens = word_tokenize(new_string, format="text")
            tokens = list(tokens.split(" "))
            if tokens != [] and tokens != [''] and len(tokens) != 1:
                # For training data, add tags
                yield gensim.models.doc2vec.TaggedDocument(tokens, [i])
                my_dict['id'].append(i)
                my_dict['file_dir'].append(test_file)
                my_dict['content'].append(line[:-1])
                i += 1

def predict_data (test_file):
    my_answer = {'id': [], 'nd_tep_input': [], 'nd_corpus' : [], 'file_dir': []}
    test_file = change_file_extension (test_file)

    #test df
    my_dict_test = {'id': [], 'file_dir': [], 'content' : []}
    test_corpus = list(read_corpus(my_dict_test, test_file))
    test_df = pd.DataFrame(data = my_dict_test)

    #load train df
    csv_file = r'E:\DATN\Duplicate_Text\app\data\train_data.csv'
    df = pd.read_csv(csv_file)

    #load model
    temp_file = datapath(r'E:\DATN\Duplicate_Text\app\doc2vec\doc2vec_model.model')
    model = gensim.models.doc2vec.Doc2Vec.load(temp_file)

    #load train_corpus
    file_name = r'E:\DATN\Duplicate_Text\app\data\train_corpus.pkl'
    open_file = open(file_name, "rb")
    train_corpus = pickle.load(open_file)
    open_file.close()

    i = 0
    for doc_id in range (len(test_corpus)):
        inferred_vector = model.infer_vector(test_corpus[doc_id].words)
        sims = model.dv.most_similar([inferred_vector], topn=len(model.dv))
        if (sims[0][1] > 0.83):
            i +=1
            id = sims[0][0]
            my_answer['id'].append(i)
            my_answer['nd_tep_input'].append(test_df.at[doc_id,'content'])
            my_answer['nd_corpus'].append(df.at[id,'content'])
            split_tup = os.path.splitext(df.at[id,'file_dir'])
            file_name = split_tup[0] + ".docx"
            my_answer['file_dir'].append(file_name)

    if i!= 0:
        i /=  len(test_df)
    duplicate_rate = {'duplicate_rate': i*100}
    i = 0
    my_answer.update(duplicate_rate)

    #duplicate_rate_files
    duplicate_rate_files = {'duplicate_rate_files':[]}
    file_dir = my_answer['file_dir']

    result = list(dict.fromkeys(my_answer['file_dir']))
    duplicate_rate_files_1 = {}
    for i in result:
        split_tup = os.path.splitext(i)
        file_name = split_tup[0] + ".txt"
        count = file_dir.count(i)
        sum = df[df['file_dir']==file_name]['id'].count()
        print (sum)
        print (df['file_dir'])
        print(i)
        duplicate_rate_files_2 = count*1.0/sum
        duplicate_rate_files_1[i] = duplicate_rate_files_2*100

    duplicate_rate_files['duplicate_rate_files'].append(duplicate_rate_files_1)
    my_answer.update(duplicate_rate_files)
    return my_answer

'''item = {'id': [], 'content_0': [], 'content_1' : [], 'file_dir': []}
path = r'E:\DATN\dataframe\test_file'
file_name = '2.txt'
test_file = os.path.join(path, file_name)
print(predict_data (test_file, item))'''
 
